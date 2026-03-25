"""PDF / DOCX report generation for incidents, daily summaries, and analytics."""

from __future__ import annotations

import logging
import os
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from sqlalchemy import select, func, and_
from sqlalchemy.ext.asyncio import AsyncSession

from backend.config import settings
from backend.database import async_session
from backend.models import Alert, Camera, Case, CaseEvidence, Event, Zone
from backend.models.models import AlertSeverity, AlertStatus, CaseStatus

logger = logging.getLogger(__name__)

# Default output directory for generated reports
REPORTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "reports",
)


class ReportGenerator:
    """Generates DOCX reports for security operations.

    All ``generate_*`` methods return the absolute file path of the
    written report.
    """

    def __init__(self, reports_dir: str = REPORTS_DIR) -> None:
        self.reports_dir = reports_dir
        os.makedirs(self.reports_dir, exist_ok=True)

    # ── Incident report ──────────────────────────────────────

    async def generate_incident_report(
        self,
        case_id: str,
    ) -> str:
        """Generate a comprehensive DOCX incident report for a case.

        Returns the absolute path to the written file.
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        async with async_session() as session:
            # Fetch case
            result = await session.execute(
                select(Case).where(Case.id == uuid.UUID(case_id))
            )
            case = result.scalar_one_or_none()
            if case is None:
                raise ValueError(f"Case not found: {case_id}")

            # Fetch evidence
            evidence_result = await session.execute(
                select(CaseEvidence)
                .where(CaseEvidence.case_id == case.id)
                .order_by(CaseEvidence.added_at)
            )
            evidence_items = evidence_result.scalars().all()

            # Fetch linked alerts (via events referenced in evidence)
            event_ids = [
                e.reference_id
                for e in evidence_items
                if e.evidence_type == "event" and e.reference_id
            ]
            linked_alerts: list = []
            if event_ids:
                alert_result = await session.execute(
                    select(Alert)
                    .where(Alert.event_id.in_(event_ids))
                    .order_by(Alert.created_at)
                )
                linked_alerts = alert_result.scalars().all()

        # Build document
        doc = Document()

        # Title
        title = doc.add_heading("SENTINEL AI -- Incident Report", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata table
        doc.add_heading("Case Details", level=1)
        table = doc.add_table(rows=7, cols=2)
        table.style = "Light Grid Accent 1"
        meta_rows = [
            ("Case ID", str(case.id)),
            ("Title", case.title or ""),
            ("Status", case.status.value if case.status else ""),
            ("Priority", case.priority.value if case.priority else ""),
            ("Created", case.created_at.strftime("%Y-%m-%d %H:%M UTC") if case.created_at else ""),
            ("Updated", case.updated_at.strftime("%Y-%m-%d %H:%M UTC") if case.updated_at else ""),
            ("Closed", case.closed_at.strftime("%Y-%m-%d %H:%M UTC") if case.closed_at else "N/A"),
        ]
        for idx, (label, value) in enumerate(meta_rows):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = value

        # Description
        if case.description:
            doc.add_heading("Description", level=1)
            doc.add_paragraph(case.description)

        # AI Insights
        if case.ai_insights:
            doc.add_heading("AI Insights", level=1)
            for key, val in case.ai_insights.items():
                doc.add_paragraph(f"{key}: {val}", style="List Bullet")

        # Summary
        if case.summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(case.summary)

        # Evidence
        if evidence_items:
            doc.add_heading("Evidence Items", level=1)
            for ev in evidence_items:
                doc.add_heading(ev.title, level=2)
                doc.add_paragraph(f"Type: {ev.evidence_type}")
                doc.add_paragraph(
                    f"Added: {ev.added_at.strftime('%Y-%m-%d %H:%M UTC') if ev.added_at else 'N/A'}"
                )
                if ev.content:
                    doc.add_paragraph(ev.content)
                if ev.file_url:
                    doc.add_paragraph(f"File: {ev.file_url}")

        # Linked alerts
        if linked_alerts:
            doc.add_heading("Linked Alerts", level=1)
            alert_table = doc.add_table(rows=1, cols=5)
            alert_table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Alert ID", "Title", "Severity", "Status", "Created"]):
                alert_table.rows[0].cells[i].text = h

            for alert in linked_alerts:
                row = alert_table.add_row()
                row.cells[0].text = str(alert.id)[:8]
                row.cells[1].text = alert.title or ""
                row.cells[2].text = alert.severity.value if alert.severity else ""
                row.cells[3].text = alert.status.value if alert.status else ""
                row.cells[4].text = (
                    alert.created_at.strftime("%Y-%m-%d %H:%M") if alert.created_at else ""
                )

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated by SENTINEL AI on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        )
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save
        filename = f"incident_report_{case_id[:8]}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.reports_dir, filename)
        doc.save(filepath)

        logger.info("Incident report generated: %s", filepath)
        return os.path.abspath(filepath)

    # ── Daily summary ────────────────────────────────────────

    async def generate_daily_summary(
        self,
        date: Optional[datetime] = None,
    ) -> str:
        """Generate a daily operations summary DOCX.

        Parameters
        ----------
        date : datetime, optional
            The day to summarise.  Defaults to today (UTC).

        Returns the absolute path to the written file.
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if date is None:
            date = datetime.now(timezone.utc)

        day_start = date.replace(hour=0, minute=0, second=0, microsecond=0)
        day_end = day_start + timedelta(days=1)

        async with async_session() as session:
            # Event count
            event_count_q = await session.execute(
                select(func.count(Event.id)).where(
                    and_(Event.timestamp >= day_start, Event.timestamp < day_end)
                )
            )
            event_count = event_count_q.scalar() or 0

            # Alerts by severity
            alert_sev_q = await session.execute(
                select(Alert.severity, func.count(Alert.id).label("cnt"))
                .where(
                    and_(Alert.created_at >= day_start, Alert.created_at < day_end)
                )
                .group_by(Alert.severity)
            )
            alert_by_sev = {
                row.severity.value if row.severity else "unknown": row.cnt
                for row in alert_sev_q.all()
            }

            # Total alerts today
            total_alerts = sum(alert_by_sev.values())

            # Camera statuses
            cam_q = await session.execute(select(Camera))
            cameras = cam_q.scalars().all()

            # Zones
            zone_q = await session.execute(
                select(Zone).where(Zone.is_active == True)  # noqa: E712
            )
            zones = zone_q.scalars().all()

            # Top event types
            event_types_q = await session.execute(
                select(Event.event_type, func.count(Event.id).label("cnt"))
                .where(
                    and_(Event.timestamp >= day_start, Event.timestamp < day_end)
                )
                .group_by(Event.event_type)
                .order_by(func.count(Event.id).desc())
                .limit(10)
            )
            top_event_types = event_types_q.all()

        # Build document
        doc = Document()

        title = doc.add_heading("SENTINEL AI -- Daily Operations Summary", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(
            f"Date: {day_start.strftime('%A, %B %d, %Y')}"
        ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Overview
        doc.add_heading("Overview", level=1)
        overview_table = doc.add_table(rows=4, cols=2)
        overview_table.style = "Light Grid Accent 1"
        overview_data = [
            ("Total Events", str(event_count)),
            ("Total Alerts", str(total_alerts)),
            ("Cameras Online", str(sum(1 for c in cameras if c.status and c.status.value == "online"))),
            ("Active Zones", str(len(zones))),
        ]
        for idx, (label, value) in enumerate(overview_data):
            overview_table.cell(idx, 0).text = label
            overview_table.cell(idx, 1).text = value

        # Alerts breakdown
        doc.add_heading("Alerts by Severity", level=1)
        sev_table = doc.add_table(rows=1, cols=2)
        sev_table.style = "Light Grid Accent 1"
        sev_table.rows[0].cells[0].text = "Severity"
        sev_table.rows[0].cells[1].text = "Count"
        for sev in ["critical", "high", "medium", "low", "info"]:
            row = sev_table.add_row()
            row.cells[0].text = sev.capitalize()
            row.cells[1].text = str(alert_by_sev.get(sev, 0))

        # Top event types
        if top_event_types:
            doc.add_heading("Top Event Types", level=1)
            et_table = doc.add_table(rows=1, cols=2)
            et_table.style = "Light Grid Accent 1"
            et_table.rows[0].cells[0].text = "Event Type"
            et_table.rows[0].cells[1].text = "Count"
            for row_data in top_event_types:
                row = et_table.add_row()
                row.cells[0].text = row_data.event_type or "unknown"
                row.cells[1].text = str(row_data.cnt)

        # Zone occupancy
        if zones:
            doc.add_heading("Zone Occupancy Snapshot", level=1)
            z_table = doc.add_table(rows=1, cols=4)
            z_table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Zone", "Type", "Occupancy", "Max"]):
                z_table.rows[0].cells[i].text = h
            for z in zones:
                row = z_table.add_row()
                row.cells[0].text = z.name
                row.cells[1].text = z.zone_type or ""
                row.cells[2].text = str(z.current_occupancy or 0)
                row.cells[3].text = str(z.max_occupancy) if z.max_occupancy else "N/A"

        # Camera status
        doc.add_heading("Camera Status", level=1)
        c_table = doc.add_table(rows=1, cols=3)
        c_table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Camera", "Location", "Status"]):
            c_table.rows[0].cells[i].text = h
        for cam in cameras:
            row = c_table.add_row()
            row.cells[0].text = cam.name
            row.cells[1].text = cam.location or ""
            row.cells[2].text = cam.status.value if cam.status else "unknown"

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated by SENTINEL AI on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        )
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save
        filename = f"daily_summary_{day_start.strftime('%Y%m%d')}_{datetime.now(timezone.utc).strftime('%H%M%S')}.docx"
        filepath = os.path.join(self.reports_dir, filename)
        doc.save(filepath)

        logger.info("Daily summary report generated: %s", filepath)
        return os.path.abspath(filepath)

    # ── Analytics report ─────────────────────────────────────

    async def generate_analytics_report(
        self,
        params: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Generate a DOCX analytics report with configurable parameters.

        Parameters
        ----------
        params : dict, optional
            ``days`` (int, default 7) -- look-back period.
            ``title`` (str) -- custom report title.
            ``include_sections`` (list[str]) -- which sections to include;
            defaults to all of:
            ``["alerts", "events", "cameras", "zones", "threats"]``.

        Returns the absolute path to the written file.
        """
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        params = params or {}
        days = params.get("days", 7)
        report_title = params.get("title", "SENTINEL AI -- Analytics Report")
        include_sections = params.get(
            "include_sections",
            ["alerts", "events", "cameras", "zones", "threats"],
        )

        cutoff = datetime.now(timezone.utc) - timedelta(days=days)

        async with async_session() as session:
            # Alerts summary
            alert_sev_q = await session.execute(
                select(Alert.severity, func.count(Alert.id).label("cnt"))
                .where(Alert.created_at >= cutoff)
                .group_by(Alert.severity)
            )
            alert_by_sev = {
                (row.severity.value if row.severity else "unknown"): row.cnt
                for row in alert_sev_q.all()
            }

            alert_status_q = await session.execute(
                select(Alert.status, func.count(Alert.id).label("cnt"))
                .where(Alert.created_at >= cutoff)
                .group_by(Alert.status)
            )
            alert_by_status = {
                (row.status.value if row.status else "unknown"): row.cnt
                for row in alert_status_q.all()
            }

            # Events summary
            event_count_q = await session.execute(
                select(func.count(Event.id)).where(Event.timestamp >= cutoff)
            )
            event_count = event_count_q.scalar() or 0

            event_types_q = await session.execute(
                select(Event.event_type, func.count(Event.id).label("cnt"))
                .where(Event.timestamp >= cutoff)
                .group_by(Event.event_type)
                .order_by(func.count(Event.id).desc())
                .limit(15)
            )
            top_event_types = event_types_q.all()

            # Cameras
            cam_q = await session.execute(select(Camera).order_by(Camera.name))
            cameras = cam_q.scalars().all()

            # Zones
            zone_q = await session.execute(
                select(Zone).where(Zone.is_active == True).order_by(Zone.name)  # noqa: E712
            )
            zones = zone_q.scalars().all()

            # Threat types
            threat_types_q = await session.execute(
                select(Alert.threat_type, func.count(Alert.id).label("cnt"))
                .where(
                    and_(
                        Alert.created_at >= cutoff,
                        Alert.threat_type.isnot(None),
                    )
                )
                .group_by(Alert.threat_type)
                .order_by(func.count(Alert.id).desc())
                .limit(15)
            )
            top_threats = threat_types_q.all()

        # Build document
        doc = Document()
        title_heading = doc.add_heading(report_title, level=0)
        title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(
            f"Period: last {days} days  |  Generated: "
            f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Alerts section
        if "alerts" in include_sections:
            doc.add_heading("Alert Analysis", level=1)
            total_alerts = sum(alert_by_sev.values())
            doc.add_paragraph(f"Total alerts in period: {total_alerts}")

            sev_table = doc.add_table(rows=1, cols=2)
            sev_table.style = "Light Grid Accent 1"
            sev_table.rows[0].cells[0].text = "Severity"
            sev_table.rows[0].cells[1].text = "Count"
            for sev in ["critical", "high", "medium", "low", "info"]:
                row = sev_table.add_row()
                row.cells[0].text = sev.capitalize()
                row.cells[1].text = str(alert_by_sev.get(sev, 0))

            doc.add_paragraph("")
            status_table = doc.add_table(rows=1, cols=2)
            status_table.style = "Light Grid Accent 1"
            status_table.rows[0].cells[0].text = "Status"
            status_table.rows[0].cells[1].text = "Count"
            for status_val, cnt in sorted(alert_by_status.items()):
                row = status_table.add_row()
                row.cells[0].text = status_val.capitalize()
                row.cells[1].text = str(cnt)

        # Events section
        if "events" in include_sections:
            doc.add_heading("Event Analysis", level=1)
            doc.add_paragraph(f"Total events in period: {event_count}")

            if top_event_types:
                et_table = doc.add_table(rows=1, cols=2)
                et_table.style = "Light Grid Accent 1"
                et_table.rows[0].cells[0].text = "Event Type"
                et_table.rows[0].cells[1].text = "Count"
                for row_data in top_event_types:
                    row = et_table.add_row()
                    row.cells[0].text = row_data.event_type or "unknown"
                    row.cells[1].text = str(row_data.cnt)

        # Camera section
        if "cameras" in include_sections:
            doc.add_heading("Camera Overview", level=1)
            c_table = doc.add_table(rows=1, cols=4)
            c_table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Camera", "Location", "Status", "Active"]):
                c_table.rows[0].cells[i].text = h
            for cam in cameras:
                row = c_table.add_row()
                row.cells[0].text = cam.name
                row.cells[1].text = cam.location or ""
                row.cells[2].text = cam.status.value if cam.status else "unknown"
                row.cells[3].text = "Yes" if cam.is_active else "No"

        # Zone section
        if "zones" in include_sections:
            doc.add_heading("Zone Overview", level=1)
            z_table = doc.add_table(rows=1, cols=5)
            z_table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Zone", "Type", "Occupancy", "Max", "Alert"]):
                z_table.rows[0].cells[i].text = h
            for z in zones:
                row = z_table.add_row()
                row.cells[0].text = z.name
                row.cells[1].text = z.zone_type or ""
                row.cells[2].text = str(z.current_occupancy or 0)
                row.cells[3].text = str(z.max_occupancy) if z.max_occupancy else "N/A"
                row.cells[4].text = "Yes" if z.alert_on_breach else "No"

        # Threat breakdown
        if "threats" in include_sections and top_threats:
            doc.add_heading("Top Threat Types", level=1)
            tt_table = doc.add_table(rows=1, cols=2)
            tt_table.style = "Light Grid Accent 1"
            tt_table.rows[0].cells[0].text = "Threat Type"
            tt_table.rows[0].cells[1].text = "Count"
            for row_data in top_threats:
                row = tt_table.add_row()
                row.cells[0].text = row_data.threat_type or "unknown"
                row.cells[1].text = str(row_data.cnt)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated by SENTINEL AI on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        )
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save
        filename = (
            f"analytics_report_{days}d_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx"
        )
        filepath = os.path.join(self.reports_dir, filename)
        doc.save(filepath)

        logger.info("Analytics report generated: %s", filepath)
        return os.path.abspath(filepath)


# Singleton
report_generator = ReportGenerator()
